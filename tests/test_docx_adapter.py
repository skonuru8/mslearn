from mslearn.adapters.docx_src import load_docx


def _build_docx(tmp_path):
    import docx

    path = tmp_path / "doc.docx"
    d = docx.Document()
    d.add_heading("A", level=1)
    d.add_paragraph("intro")
    d.add_heading("A.1", level=2)
    d.add_paragraph("deep")
    d.save(str(path))
    return path


def test_load_docx_heading_ancestry(tmp_path):
    path = _build_docx(tmp_path)

    doc = load_docx(path, role="supplement")

    assert doc.source_type == "docx" and doc.role == "supplement"
    by_text = {u.text.strip(): u.section_path for u in doc.units}
    assert by_text["intro"] == ("A",)
    assert by_text["deep"] == ("A", "A.1")


def test_load_docx_no_headings_is_flat(tmp_path):
    import docx

    path = tmp_path / "flat.docx"
    d = docx.Document()
    d.add_paragraph("just body text")
    d.save(str(path))

    doc = load_docx(path)

    assert len(doc.units) == 1
    assert doc.units[0].section_path == ()
