import re
from pathlib import Path

from mslearn.adapters.base import Locator, SourceDocument, StructuralUnit, make_source_id

_HEADING_RE = re.compile(r"^Heading (\d+)$")


def _heading_level(style_name: str) -> int | None:
    if style_name == "Title":
        return 1
    match = _HEADING_RE.match(style_name or "")
    return int(match.group(1)) if match else None


def load_docx(path: Path | str, role: str = "supplement") -> SourceDocument:
    import docx  # heavy import: only pulled in when a .docx is actually loaded

    path = Path(path)
    document = docx.Document(str(path))

    units: list[StructuralUnit] = []
    stack: list[tuple[int, str]] = []
    buf: list[str] = []

    def flush() -> None:
        text = "\n".join(buf).strip()
        buf.clear()
        if not text:
            return
        units.append(
            StructuralUnit(
                index=len(units), title="", text=text,
                locator=Locator(kind="para", para_index=len(units)),
                section_path=tuple(title for _, title in stack),
            )
        )

    for para in document.paragraphs:
        level = _heading_level(para.style.name if para.style else "")
        text = para.text.strip()
        if level is not None:
            flush()
            if text:
                while stack and stack[-1][0] >= level:
                    stack.pop()
                stack.append((level, text))
        elif text:
            buf.append(text)
    flush()

    return SourceDocument(
        source_id=make_source_id(str(path)), source_type="docx",
        role=role, title=path.stem, units=units,
    )
